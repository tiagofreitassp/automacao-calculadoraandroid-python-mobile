import os, shutil
from time import sleep
from appium import webdriver
from appium.options.common import AppiumOptions
from docx import Document
from docx.shared import Inches

class emuladorDriver():
    dir = '../evidencias/'

    def criarPastaEvidencia(self, nPasta):
            d = nPasta
            if os.path.exists(d) == True:
                print("Diretório já existe")
                #os.rmdir(d)
                shutil.rmtree(d)
                print('Diretório apagado')
                os.makedirs(d)
                print('Diretório recriado')
            else:
                print("Diretório não existe")
                os.makedirs(d)
                print('Diretório criado')

    def gerarScreenshot(self, nPasta, nEvidencia):
        sleep(1)
        self.driver.get_screenshot_as_file(nPasta + "/" + nEvidencia + ".png")

    def criarDocumentoDeEvidencia(self, diretorioEvidencia,id ,nomeEvidencia):
        try:
            document = Document()

            document.add_heading('Evidências: Calculadora Android', 0)
            p = document.add_paragraph(nomeEvidencia)

            document.add_paragraph(id+'_Tela01')
            document.add_picture(diretorioEvidencia + '/Ev1.png', width=Inches(4.14))
            #document.add_page_break()

            document.add_paragraph(id+'_Tela02')
            document.add_picture(diretorioEvidencia + '/Ev2.png', width=Inches(4.14))
            #document.add_page_break()

            document.save(diretorioEvidencia + '/' + nomeEvidencia + '.docx')
            print("Documento com as evidencias gerada com sucesso!")
        except FileExistsError as error:
            print(error)
        except FileNotFoundError as error:
            print(error)
        except BufferError as error:
            print(error)
        except BufferError as error:
            print(error)
        except:
            print("Não foi possivel criar o documento com as evidencias!")

    def criarDriver(self):
        desired_caps = {
            "platformName": "Android",
            "deviceName": "emulator-5554",
            "automationName": "uiautomator2",
            "appPackage": "com.android.calculator2",
            "appActivity": "com.android.calculator2.Calculator"
        }
        option = AppiumOptions().load_capabilities(desired_caps)
        self.driver = webdriver.Remote(command_executor=f"http://127.0.0.1:4723", options=option)
        self.driver.implicitly_wait(15)

    def fecharDriver(self):
        self.driver.quit()